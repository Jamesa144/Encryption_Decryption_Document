import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from cryptography.fernet import Fernet
import os
import base64

# Function to generate a key for encryption
def generate_key():
    return Fernet.generate_key()

# Function to encrypt text
def encrypt_text(text, key):
    cipher = Fernet(key)
    return base64.b64encode(cipher.encrypt(text.encode())).decode()

# Function to decrypt text
def decrypt_text(encrypted_text, key):
    cipher = Fernet(key)
    return cipher.decrypt(base64.b64decode(encrypted_text)).decode()

# Function to read and encrypt the text content of a Word document
def read_and_encrypt_docx(file_path, key):
    doc = Document(file_path)
    encrypted_doc = Document()
    
    for paragraph in doc.paragraphs:
        encrypted_text = encrypt_text(paragraph.text, key)
        new_paragraph = encrypted_doc.add_paragraph()
        new_paragraph.text = encrypted_text
        new_paragraph.style = paragraph.style
        
        for run in paragraph.runs:
            new_run = new_paragraph.add_run()
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
    
    return encrypted_doc

# Modified function to read and decrypt the text content of a Word document
def read_and_decrypt_docx(file_path, key):
    doc = Document(file_path)
    decrypted_doc = Document()
    
    for paragraph in doc.paragraphs:
        try:
            decrypted_text = decrypt_text(paragraph.text, key)
        except:
            # If decryption fails, keep the original text
            decrypted_text = paragraph.text
        
        new_paragraph = decrypted_doc.add_paragraph()
        new_paragraph.style = paragraph.style
        
        new_run = new_paragraph.add_run(decrypted_text)
        
        # Copy formatting from the first run of the original paragraph
        if paragraph.runs:
            original_run = paragraph.runs[0]
            new_run.bold = original_run.bold
            new_run.italic = original_run.italic
            new_run.underline = original_run.underline
            new_run.font.name = original_run.font.name
            new_run.font.size = original_run.font.size
            if original_run.font.color.rgb:
                new_run.font.color.rgb = original_run.font.color.rgb
    
    return decrypted_doc


# Function to save the encryption key to a file
def save_key(key, file_path):
    key_file_path = f"{os.path.splitext(file_path)[0]}_EncryptKey.txt"
    with open(key_file_path, 'wb') as key_file:
        key_file.write(key)
    return key_file_path

# Function to load the encryption key from a file
def load_key(file_path):
    with open(file_path, 'rb') as key_file:
        return key_file.read()

# Function to handle file selection and encryption
def process_files_encrypt():
    input_file = filedialog.askopenfilename(
        title="Select Word Document to Encrypt",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not input_file:
        messagebox.showwarning("Warning", "No input file selected")
        return
    
    output_file = filedialog.asksaveasfilename(
        title="Save Encrypted Word Document",
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not output_file:
        messagebox.showwarning("Warning", "No output file selected")
        return
    
    key = generate_key()
    key_file_path = save_key(key, input_file)
    print(f"Encryption key saved to: {key_file_path}")

    try:
        encrypted_doc = read_and_encrypt_docx(input_file, key)
        encrypted_doc.save(output_file)
        messagebox.showinfo("Success", "Document encrypted and saved successfully")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Function to handle file selection and decryption
def process_files_decrypt():
    input_file = filedialog.askopenfilename(
        title="Select Encrypted Word Document to Decrypt",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not input_file:
        messagebox.showwarning("Warning", "No input file selected")
        return
    
    key_file = filedialog.askopenfilename(
        title="Select Encryption Key File",
        filetypes=[("Text Files", "*.txt")]
    )
    if not key_file:
        messagebox.showwarning("Warning", "No key file selected")
        return
    
    output_file = filedialog.asksaveasfilename(
        title="Save Decrypted Word Document",
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not output_file:
        messagebox.showwarning("Warning", "No output file selected")
        return
    
    try:
        key = load_key(key_file)
        decrypted_doc = read_and_decrypt_docx(input_file, key)
        decrypted_doc.save(output_file)
        messagebox.showinfo("Success", "Document decrypted and saved successfully")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Create the main GUI window
root = tk.Tk()
root.title("Word Document Encryptor/Decryptor")

# Set the size of the window
root.geometry("450x200")

# Create buttons to start the encryption and decryption processes
encrypt_button = tk.Button(root, text="Encrypt Word Document", command=process_files_encrypt)
encrypt_button.pack(padx=20, pady=10)

decrypt_button = tk.Button(root, text="Decrypt Word Document", command=process_files_decrypt)
decrypt_button.pack(padx=20, pady=10)

# Run the GUI event loop
root.mainloop()